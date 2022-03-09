import pandas as pd
from docx import Document
import os.path
import json
from docx.shared import Inches

epoch = range(20,61,10)
root_folder = r"C:\Users\Christina\Documents\DIP\Week9\samples\flickr_original"
#val_folder = r"C:\Users\Christina\Documents\DIP\Week4\img\test"
val_folder = r"C:\Users\Christina\Documents\DIP\Week3\flickr8k\images"
document = Document()

for ep in epoch:

    filename = os.path.join(root_folder, "sample_epoch_{}.xlsx".format(ep))
    df = pd.read_csv(filename, header=0)
    document.add_heading("Epoch: " + str(ep), level = 2)

    num_rows = len(df)
    table = document.add_table(rows=num_rows, cols=2)

    for idx, rows in df.iterrows():
        pic_cell = table.cell(idx, 0)
        paragraph = pic_cell.paragraphs[0]
        run = paragraph.add_run()
        #run.add_picture(os.path.join(val_folder, rows["image_id"]+".jpg"), width = Inches(2.0))
        run.add_picture(os.path.join(val_folder, rows["image"]), width = Inches(2.0))

        text_cell = table.cell(idx, 1)
        text_cell.text = rows["pred_caption"]

    document.add_page_break()

document.save(r'C:\Users\Christina\Documents\DIP\Week9\samples\flickr_original.docx')

print("Done!")

